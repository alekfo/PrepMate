"""PDF and DOCX export for Resume."""
import io
import os

_FONTS_DIR = os.path.join(os.path.dirname(__file__), 'fonts')
_FONT_REG = os.path.join(_FONTS_DIR, 'DejaVuSans.ttf')
_FONT_BOLD = os.path.join(_FONTS_DIR, 'DejaVuSans-Bold.ttf')

# Colors (R, G, B)
_C_TEXT = (26, 26, 26)
_C_SUB = (75, 75, 75)
_C_MUTED = (130, 130, 130)
_C_LINE = (210, 210, 210)


# ── PDF ──────────────────────────────────────────────────────────────────────

def generate_pdf(resume, sections, experience_label: str) -> bytes:
    from fpdf import FPDF

    PAGE_W = 210
    M = 18          # margin
    CW = PAGE_W - 2 * M  # content width

    pdf = FPDF(format='A4')
    pdf.set_margins(M, M, M)
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()

    pdf.add_font('dv', '', _FONT_REG)
    pdf.add_font('dv', 'B', _FONT_BOLD)

    contacts = sections.get('contacts')
    c = contacts.display_content if contacts else {}
    full_name = c.get('full_name', '')

    # ── Header ──
    PHOTO_W = 45
    PHOTO_H = 45
    has_photo = bool(resume.photo)
    text_w = CW - PHOTO_W - 7 if has_photo else CW
    photo_x = M + text_w + 7
    photo_y = M

    # Name — shrink font until it fits on one line; wrap as fallback with 'L' align
    name_size = 21
    pdf.set_font('dv', 'B', name_size)
    while name_size > 14 and pdf.get_string_width(full_name) > text_w - 1:
        name_size -= 1
        pdf.set_font('dv', 'B', name_size)
    pdf.set_text_color(*_C_TEXT)
    pdf.set_xy(M, M)
    pdf.multi_cell(text_w, round(name_size * 0.43 + 2, 1), full_name, align='L')

    # Subtitle
    pdf.set_font('dv', '', 11)
    pdf.set_text_color(*_C_SUB)
    pdf.set_x(M)
    pdf.cell(text_w, 6, f'{resume.profession} — стаж {experience_label}')
    pdf.ln(8)

    # Contacts
    parts = [v for k, v in [
        ('email', c.get('email', '')), ('phone', c.get('phone', '')),
        ('city', c.get('city', '')), ('linkedin', c.get('linkedin', '')),
        ('github', c.get('github', '')),
    ] if v]
    if parts:
        pdf.set_font('dv', '', 8.5)
        pdf.set_text_color(*_C_MUTED)
        pdf.set_x(M)
        pdf.multi_cell(text_w, 4.5, '   ·   '.join(parts))

    header_bottom = pdf.get_y()

    # Photo
    if has_photo:
        try:
            pdf.image(resume.photo.path, x=photo_x, y=photo_y, w=PHOTO_W, h=PHOTO_H)
        except Exception:
            pass
        header_bottom = max(header_bottom, photo_y + PHOTO_H)

    pdf.set_draw_color(*_C_LINE)
    pdf.set_y(header_bottom + 10)

    # ── Section helpers ──
    def sec_title(title):
        pdf.ln(6)  # breathing room before each section
        pdf.set_font('dv', 'B', 13)
        pdf.set_text_color(*_C_TEXT)
        pdf.set_x(M)
        pdf.cell(CW, 8, title)
        pdf.ln(8)
        # solid dark underline
        y = pdf.get_y()
        pdf.set_draw_color(*_C_TEXT)
        pdf.set_line_width(0.5)
        pdf.line(M, y, PAGE_W - M, y)
        pdf.set_line_width(0.2)
        pdf.set_draw_color(*_C_LINE)
        pdf.set_y(y + 6)
        pdf.set_text_color(*_C_TEXT)

    def body(txt, size=10, color=None):
        pdf.set_font('dv', '', size)
        pdf.set_text_color(*(color or _C_TEXT))
        pdf.set_x(M)
        pdf.multi_cell(CW, 5.5, txt)

    def row_with_right(left_txt, right_txt, left_bold=False, size=10):
        """Single-line row: left text (bold opt) + right text right-aligned."""
        pdf.set_font('dv', 'B' if left_bold else '', size)
        pdf.set_text_color(*_C_TEXT)
        rw = pdf.get_string_width(right_txt) + 1
        lw = CW - rw
        pdf.set_x(M)
        pdf.cell(lw, 6, left_txt, ln=False)
        pdf.set_font('dv', '', size - 1)
        pdf.set_text_color(*_C_MUTED)
        pdf.cell(rw, 6, right_txt, align='R', ln=True)

    # ── О себе ──
    s = sections.get('summary')
    if s:
        text = s.display_content.get('text', '')
        if text:
            sec_title('О себе')
            body(text, color=_C_SUB)

    # ── Опыт работы ──
    exp = sections.get('experience')
    if exp:
        items = exp.display_content
        if isinstance(items, list) and items:
            sec_title('Опыт работы')
            for i, item in enumerate(items):
                period = f'{item.get("period_start", "")} — {item.get("period_end", "")}'
                row_with_right(item.get('position', ''), period, left_bold=True, size=10.5)

                pdf.set_font('dv', '', 9.5)
                pdf.set_text_color(*_C_SUB)
                pdf.set_x(M)
                pdf.cell(CW, 5, item.get('company', ''), ln=True)

                if item.get('responsibilities'):
                    pdf.ln(2)
                    body(item['responsibilities'], size=9.5, color=_C_TEXT)
                if item.get('achievements'):
                    body(item['achievements'], size=9.5, color=_C_SUB)
                if i < len(items) - 1:
                    # divider between experience items
                    pdf.ln(4)
                    y2 = pdf.get_y()
                    pdf.set_draw_color(*_C_LINE)
                    pdf.line(M, y2, PAGE_W - M, y2)
                    pdf.set_y(y2 + 4)

    # ── Образование ──
    edu = sections.get('education')
    if edu:
        items = edu.display_content
        if isinstance(items, list) and items:
            sec_title('Образование')
            for item in items:
                row_with_right(item.get('institution', ''), item.get('year', ''), left_bold=True)
                sub = item.get('specialty', '')
                if item.get('degree'):
                    sub += f' · {item["degree"]}'
                if sub:
                    pdf.set_font('dv', '', 9.5)
                    pdf.set_text_color(*_C_SUB)
                    pdf.set_x(M)
                    pdf.cell(CW, 5, sub, ln=True)
                pdf.ln(4)

    # ── Навыки ──
    sk = sections.get('skills')
    if sk:
        s = sk.display_content
        if s.get('hard_skills') or s.get('soft_skills'):
            sec_title('Навыки')
            if s.get('hard_skills'):
                pdf.set_font('dv', 'B', 9.5)
                pdf.set_text_color(*_C_TEXT)
                pdf.set_x(M)
                pdf.cell(CW, 5.5, 'Профессиональные навыки')
                pdf.ln(5.5)
                body(s['hard_skills'], size=9.5, color=_C_SUB)
                pdf.ln(4)
            if s.get('soft_skills'):
                pdf.set_font('dv', 'B', 9.5)
                pdf.set_text_color(*_C_TEXT)
                pdf.set_x(M)
                pdf.cell(CW, 5.5, 'Личные качества')
                pdf.ln(5.5)
                body(s['soft_skills'], size=9.5, color=_C_SUB)

    # ── Языки ──
    lang = sections.get('languages')
    if lang:
        items = lang.display_content
        if isinstance(items, list) and items:
            sec_title('Языки')
            lang_str = '   ·   '.join(
                f'{i.get("language", "")} — {i.get("level", "")}' for i in items
            )
            body(lang_str, size=10, color=_C_SUB)

    # ── Курсы и сертификаты ──
    cert = sections.get('certifications')
    if cert:
        items = cert.display_content
        real = [i for i in (items or []) if i.get('name')]
        if real:
            sec_title('Курсы и сертификаты')
            for item in real:
                # Name wraps if long
                pdf.set_font('dv', '', 10)
                pdf.set_text_color(*_C_TEXT)
                pdf.set_x(M)
                pdf.multi_cell(CW, 5.5, item.get('name', ''))
                # Platform · Year always on a separate line below
                meta_parts = [p for p in [item.get('platform', ''), item.get('year', '')] if p]
                if meta_parts:
                    pdf.set_font('dv', '', 9)
                    pdf.set_text_color(*_C_MUTED)
                    pdf.set_x(M)
                    pdf.cell(CW, 4.5, ' · '.join(meta_parts), ln=True)
                pdf.ln(4)

    return bytes(pdf.output())


# ── DOCX ─────────────────────────────────────────────────────────────────────

def generate_docx(resume, sections, experience_label: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    # Remove default paragraph spacing
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    contacts = sections.get('contacts')
    c = contacts.display_content if contacts else {}
    full_name = c.get('full_name', '')

    def _rgb(hex_str):
        h = hex_str.lstrip('#')
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def _no_space(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    def _para_border_bottom(p, color='DDDDDD'):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _table_no_borders(table):
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        bdr = OxmlElement('w:tblBorders')
        for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{name}')
            el.set(qn('w:val'), 'none')
            bdr.append(el)
        tblPr.append(bdr)

    # ── Header table: name/info left, photo right ──
    has_photo = bool(resume.photo)
    tbl = doc.add_table(rows=1, cols=2 if has_photo else 1)
    _table_no_borders(tbl)
    tbl.autofit = False

    if has_photo:
        tbl.columns[0].width = Cm(12.5)
        tbl.columns[1].width = Cm(4.0)
    else:
        tbl.columns[0].width = Cm(16.5)

    left = tbl.cell(0, 0)

    # Name
    p_name = left.add_paragraph(full_name)
    _no_space(p_name)
    r = p_name.runs[0]
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = _rgb('1a1a1a')

    # Subtitle
    p_sub = left.add_paragraph(f'{resume.profession} — стаж {experience_label}')
    p_sub.paragraph_format.space_before = Pt(3)
    _no_space(p_sub)
    p_sub.paragraph_format.space_after = Pt(2)
    r2 = p_sub.runs[0]
    r2.font.size = Pt(11)
    r2.font.color.rgb = _rgb('4b4b4b')

    # Contacts
    contact_parts = [v for k, v in [
        ('email', c.get('email', '')), ('phone', c.get('phone', '')),
        ('city', c.get('city', '')), ('linkedin', c.get('linkedin', '')),
        ('github', c.get('github', '')),
    ] if v]
    if contact_parts:
        p_ct = left.add_paragraph('   ·   '.join(contact_parts))
        _no_space(p_ct)
        p_ct.paragraph_format.space_before = Pt(4)
        r3 = p_ct.runs[0]
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = _rgb('888888')

    # Photo cell
    if has_photo:
        right = tbl.cell(0, 1)
        p_ph = right.paragraphs[0]
        p_ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _no_space(p_ph)
        run = p_ph.add_run()
        try:
            run.add_picture(resume.photo.path, width=Cm(3.5))
        except Exception:
            pass

    # Separator paragraph — dark line after header
    p_sep = doc.add_paragraph()
    _no_space(p_sep)
    p_sep.paragraph_format.space_before = Pt(10)
    _para_border_bottom(p_sep, color='1a1a1a')

    def sec_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        rh = p.add_run(title)
        rh.font.size = Pt(13)
        rh.font.bold = True
        rh.font.color.rgb = _rgb('1a1a1a')
        _para_border_bottom(p, color='1a1a1a')

    def body_para(txt, size=10, bold=False, color='333333', before=0, after=2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(txt)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = _rgb(color)
        return p

    def row_bold_right(left_txt, right_txt, size=10.5):
        """Paragraph with bold left text and muted right text."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        rl = p.add_run(left_txt)
        rl.font.size = Pt(size)
        rl.font.bold = True
        rl.font.color.rgb = _rgb('1a1a1a')
        if right_txt:
            p.add_run('  ')
            rr = p.add_run(right_txt)
            rr.font.size = Pt(9)
            rr.font.bold = False
            rr.font.color.rgb = _rgb('999999')

    # ── О себе ──
    s = sections.get('summary')
    if s:
        text = s.display_content.get('text', '')
        if text:
            sec_heading('О себе')
            body_para(text, size=10, color='4b4b4b')

    # ── Опыт работы ──
    exp = sections.get('experience')
    if exp:
        items = exp.display_content
        if isinstance(items, list) and items:
            sec_heading('Опыт работы')
            for item in items:
                period = f'{item.get("period_start", "")} — {item.get("period_end", "")}'
                row_bold_right(item.get('position', ''), period)
                body_para(item.get('company', ''), size=10, color='555555', before=0, after=2)
                if item.get('responsibilities'):
                    body_para(item['responsibilities'], size=9.5, color='333333', before=2, after=2)
                if item.get('achievements'):
                    body_para(item['achievements'], size=9.5, color='555555', before=0, after=4)

    # ── Образование ──
    edu = sections.get('education')
    if edu:
        items = edu.display_content
        if isinstance(items, list) and items:
            sec_heading('Образование')
            for item in items:
                row_bold_right(item.get('institution', ''), item.get('year', ''))
                sub = item.get('specialty', '')
                if item.get('degree'):
                    sub += f' · {item["degree"]}'
                if sub:
                    body_para(sub, size=9.5, color='555555', before=1, after=4)

    # ── Навыки ──
    sk = sections.get('skills')
    if sk:
        s = sk.display_content
        if s.get('hard_skills') or s.get('soft_skills'):
            sec_heading('Навыки')
            if s.get('hard_skills'):
                body_para('Профессиональные навыки', size=9.5, bold=True, color='1a1a1a', before=2, after=1)
                body_para(s['hard_skills'], size=9.5, color='4b4b4b', before=0, after=4)
            if s.get('soft_skills'):
                body_para('Личные качества', size=9.5, bold=True, color='1a1a1a', before=0, after=1)
                body_para(s['soft_skills'], size=9.5, color='4b4b4b', before=0, after=4)

    # ── Языки ──
    lang = sections.get('languages')
    if lang:
        items = lang.display_content
        if isinstance(items, list) and items:
            sec_heading('Языки')
            lang_str = '   ·   '.join(
                f'{i.get("language", "")} — {i.get("level", "")}' for i in items
            )
            body_para(lang_str, size=10, color='4b4b4b')

    # ── Курсы и сертификаты ──
    cert = sections.get('certifications')
    if cert:
        items = cert.display_content
        real = [i for i in (items or []) if i.get('name')]
        if real:
            sec_heading('Курсы и сертификаты')
            for item in real:
                row_bold_right(item.get('name', ''), item.get('year', ''), size=10)
                if item.get('platform'):
                    body_para(item['platform'], size=9, color='888888', before=0, after=3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()